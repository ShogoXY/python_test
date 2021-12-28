# bedziemy próbować zrobić działający program na plikach excel za pomocą Pandas
import pandas as pd
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import numpy as np
from numpy import nan
import openpyxl
from pandastable import *
from tkinter import *
import re
# from openpyxl import load_workbook

import time
from re import search
from tkinter.simpledialog import askstring
from tkinter.messagebox import showinfo

main_window = Tk()
main_window.title("Program do rozliczeń")
main_window.state('zoomed')
e_klient = askstring(0,"Podaj dla jakiego Klienta")




# file path
# path_excel_open = "C:\\Users\\Dariusz\\github\\python_test\\serwi.xlsx"
path_excel_open = "\\\\192.168.2.150\\serwis\\SERWISOWE\\Test\\SERWISY2.xlsx"
path_docx = "C:\\Users\\Dariusz\\github\\python_test\\Rozliczenie serwisowe.docx"
path_excel_save = "C:\\Users\\Dariusz\\github\\python_test\\serwi2.xlsx"
path_docx2 = "C:\\Users\\Dariusz\\github\\python_test\\rozliczenie"

doc = docx.Document(path_docx)
doc = DocxTemplate(path_docx)

ct = (time.strftime('%d.%m.%Y'))
cr = time.strftime('%y%m%d%H%M')
cr2 = ("CRS"+cr)
my_label = Label(main_window, text=ct)
my_label2 = Label(main_window, text=cr2)
my_label.pack()
my_label2.pack()



#e_klient = askstring(main_window, width=70, borderwidth=2)
#e_klient.insert(0, "Podaj dla jakiego Klienta")

l_klient = Label(main_window, text="Rozliczenie dla klienta - " + format(e_klient))
l_klient.pack()
#e_klient.pack()


# Usuwanie tekstu (temporaty text)
def temp_text(e):
   e_klient.delete(0,"end")
    
#e_klient.bind("<FocusIn>", temp_text)

#######################################


def press_klient(event):
    global client
    client = e_klient.get()
    # e_klient.delete(0, "end")
    e_klient.insert(0, '')
    my_label3 = Label(main_window, text=client)
    my_label3.place(x=20, y=60)
    my_label3.pack()
    pt = Table(frame, dataframe=df)
    pt.show()


def click_klient():
    global client
    client=e_klient.get()
    # e_klient.delete(0, "end")
    e_klient.insert(0, '')
    my_label3 = Label(main_window, text=client)
    my_label3.place(x=20, y=60)
    my_label3.pack()
    pt = Table(frame, dataframe=df)
    pt.show()


# e_klient.bind('<Return>', press_klient)
# button1 = Button(main_window, text="klient", command=click_klient) 
# button1.pack()




# nazwa_klienta = input("podaj klienta \n")


# doc = DocxTemplate(path_docx)
# context = {'date': ct, 'cr_number': cr2, 'nazwa': nazwa_klienta, 'data': ct}
# doc.render(context)

#### wyswietl tabele 

df = pd.read_excel(path_excel_open, sheet_name='_____SERWISY______')
df = df.astype(str)
df = df.fillna('')


frame = Frame(main_window)
frame = LabelFrame(main_window, text="Tabela ze wszystkimi Numerami RMA")
frame.pack_propagate(False)
frame.pack(fill='both')
pt = Table(frame, dataframe=df)
pt.show()



# pt = Table(frame, dataframe=df)
# pt.show()



df_name = "arkusz_test"
excel_book = openpyxl.load_workbook(path_excel_save)

if df_name not in excel_book.sheetnames:
    excel_book.create_sheet(df_name)
excel_book.save(path_excel_save)

writer = pd.ExcelWriter(path_excel_save, engine='openpyxl',
                        mode='a', if_sheet_exists='replace')
writer.book = excel_book
writer.sheets = dict((ws.title, ws) for ws in excel_book.worksheets)




def search_value(keyword, df):
    search_value = '|'.join(keyword)
    searched = df[df['Unnamed: 0'].str.contains(search_value, na=False)]
    return searched


l_rozliczenie = Label(main_window, text="Poniżej proszę wprowadzić numer RMA do dodania w rozliczeniu")
l_rozliczenie.pack()
e_rma = Entry(main_window, width=70, borderwidth=2)
e_rma.insert(0, "Podaj numer RMA")
e_rma.pack()


def print_rma(event):

    if e_rma.get() == "":
        messagebox.showerror(title="Error", message="Brak wartości")
    else:       
        search_df = search_value([e_rma.get()], df)
        
        df1 = pd.DataFrame(search_df, columns=['Unnamed: 0', 'Unnamed: 5',
                        'Unnamed: 6', 'Unnamed: 8', 'Unnamed: 11'])
        df1 = df1.replace('nan', '')
        writer.save()
        df2 = pd.read_excel(path_excel_save, sheet_name='arkusz_test')
        df3 = df2.append(df1, ignore_index=True)
        df2 = df2.fillna('')
        df3.to_excel(writer, sheet_name='arkusz_test', index=False)

        df3 = df3.fillna('')
        df3.dropna(subset = ["Unnamed: 5"], inplace=True)

        pt_rma = Table(frame_rma, dataframe=df3)
        pt_rma.show()
        
        writer.save()

        
        
    for i in range(df3.shape[0]):
        doc.tables[0].add_row()
        for j in range(df3.shape[-1]):
            table2 = doc.tables[0]
            table2.cell(i+1, j+1).text = str(df3.values[i, j])
            table2.cell(i+1, 0).text = str(i+1)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    

def print_rma_click():
    
    if e_rma.get() == "":
        messagebox.showerror(title="Error", message="Brak wartości")
    else:       
        search_df = search_value([e_rma.get()], df)
        
        df1 = pd.DataFrame(search_df, columns=['Unnamed: 0', 'Unnamed: 5',
                        'Unnamed: 6', 'Unnamed: 8', 'Unnamed: 11'])
        df1 = df1.replace('nan', '')
        writer.save()
        df2 = pd.read_excel(path_excel_save, sheet_name='arkusz_test')
        df3 = df2.append(df1, ignore_index=True)
        df2 = df2.fillna('')
        df3.to_excel(writer, sheet_name='arkusz_test', index=False)

        df3 = df3.fillna('')
        df3.dropna(subset = ["Unnamed: 5"], inplace=True)

        pt_rma = Table(frame_rma, dataframe=df3)
        pt_rma.show()
        
        writer.save()
       


    for i in range(df3.shape[0]):
        doc.tables[0].add_row()
        
        for j in range(df3.shape[-1]):
            table2 = doc.tables[0]
            table2.cell(i+1, j+1).text = str(df3.values[i, j])
            table2.cell(i+1, 0).text = str(i+1)
            table2.paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
         
    

client2 = format(e_klient)

def write_to_doc():
    context = {'date': ct, 'cr_number': cr2, 'nazwa': client2, 'data': ct}
    doc.render(context)
    doc.save(path_docx2 + "\\rozliczenie_" + cr2 + "_" + client2 +".docx")
    del excel_book[df_name]
    excel_book.save(path_excel_save)
    excel_book.close()
    messagebox.showinfo("zapisano", "plik zapisany w " + path_docx2 + "\\rozliczenie_" + cr2 + "_" + client2 +".docx")
    os.system("start " + path_docx2 + "\\rozliczenie_" + cr2 + "_" + client2 +".docx") 
    main_window.destroy()


def temp_text(e):
   e_rma.delete(0,"end")

e_rma.bind("<FocusIn>", temp_text)


e_rma.bind('<Return>', print_rma)


button_rma_show = Button(main_window, text="Dodaj numer RMA", command=print_rma_click) 
button_rma_show.pack()


frame_rma = Frame(main_window)
frame_rma = LabelFrame(main_window, text="Numery RMA do wpisania do dokumentu")
frame_rma.pack_propagate(False)
frame_rma.pack(fill='both',)


button_save_doc = Button(main_window, text="Zapisz i zamknij", command=write_to_doc) 
button_save_doc.pack()



main_window.mainloop()