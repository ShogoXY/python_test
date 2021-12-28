# bedziemy próbować zrobić działający program na plikach excel za pomocą Pandas
import pandas as pd
import docx
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import numpy as np
from openpyxl import load_workbook

# file path
path_file = "C:\\Users\\Dariusz\\github\\python_test\\test.xlsx"
path_docx = "C:\\Users\\Dariusz\\github\\python_test\\word1.docx"

doc = docx.Document(path_docx)

# search value in excel sheet (if yes, pritn it)
df = pd.read_excel(path_file, sheet_name='Produkcja')
search_word = input("podaj numer seryjny do wyszukania \n")
search_df = df.loc[df['NUMER SERYJNY'] == search_word]

# print value in specific column and search with this value
komentarz_search = search_df['KOMENTARZ'].values[0]
print(komentarz_search)
komentarz_df = df.loc[df['KOMENTARZ'] == komentarz_search]
print(komentarz_df)

# write value to new sheet in same workbook
excel_book = load_workbook(path_file)
writer = pd.ExcelWriter(path_file, engine='openpyxl',
                        mode='a', if_sheet_exists='replace')
writer.book = excel_book
writer.sheets = dict((ws.title, ws) for ws in excel_book.worksheets)
df1 = pd.DataFrame(komentarz_df, columns=['NAZWA', 'NUMER SERYJNY'])
df1.to_excel(writer, sheet_name='arkusz')

# save nd exit excel
writer.save()
writer.close()


# word create table in word
table = doc.add_table(df1.shape[0]+1, df1.shape[1], style='Table Grid')


for j in range(df1.shape[-1]):
    table.cell(0, j).text = df1.columns[j]

for i in range(df1.shape[0]):
    for j in range(df1.shape[-1]):
        table.cell(i+1, j).text = str(df1.values[i, j])
        
n=(i+1)
for k in range(1,n):
    m = str(k) 
    print("test_" + m +"")



print(str(df1.values[i, j]))
doc.save(path_docx)

print(table.cell.text)

# '
#input = Document('OutputDoc.docx')

#paragraphs = []
# for para in input.paragraphs:
#    p = para.text
#    paragraphs.append(p)

#output = Document()
# for item in paragraphs:
#    output.add_paragraph(item)
# output.save('word1.docx')
#

# Imports

input_doc = Document('OutputDoc.docx')
output_doc = Document()

# Call the function


def get_para_data(output_doc_name, paragraph):

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        output_run.style.name = run.style.name
        output_run.font.size = run.font.size
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
    output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    output_para.paragraph_format.right_indent = paragraph.paragraph_format.left_indent
    output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.left_indent


for para in input_doc.paragraphs:
    get_para_data(output_doc, para)

output_doc.save('OutputDoc2.docx')
