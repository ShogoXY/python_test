# bedziemy próbować zrobić działający program na plikach excel za pomocą Pandas
import pandas as pd
import docx
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import numpy as np
from numpy import nan
import openpyxl
import re
from tkinter import *
# from openpyxl import load_workbook

import datetime
import time
from re import search
import keyboard

# ################## GUI ##################
main_window = Tk()
main_window.title("Program do rozliczeń")
main_window.geometry("500x500")



# file path
path_excel_open = "C:\\Users\\Dariusz\\github\\python_test\\serwi.xlsx"
path_docx = "C:\\Users\\Dariusz\\github\\python_test\\Rozliczenie serwisowe.docx"
path_excel_save = "C:\\Users\\Dariusz\\github\\python_test\\serwi2.xlsx"


doc = docx.Document(path_docx)


ct = (time.strftime('%d.%m.%Y'))
print(ct)
cr = time.strftime('%y%m%d%H%M')
cr2 = ("CRS"+cr)
print(cr2)

nazwa = input("podaj klienta \n")


doc = DocxTemplate(path_docx)
context = {'date': ct, 'cr_number': cr2, 'nazwa': nazwa, 'data': ct}
doc.render(context)


df = pd.read_excel(path_excel_open, sheet_name='Arkusz1')
df = df.astype(str)
df = df.fillna('')
print(df)

df_name = "arkusz_test"
excel_book = openpyxl.load_workbook(path_excel_save)

if df_name not in excel_book.sheetnames:
    excel_book.create_sheet(df_name)
excel_book.save(path_excel_save)

writer = pd.ExcelWriter(path_excel_save, engine='openpyxl',
                        mode='a', if_sheet_exists='replace')
writer.book = excel_book
writer.sheets = dict((ws.title, ws) for ws in excel_book.worksheets)

while 1:

    def search_value(keyword, df):
        search_value = '|'.join(keyword)
        searched = df[df['RMA'].str.contains(search_value, na=False)]
        return searched

    # podaj wartość do wpisania
    search_word = input("podaj numer RMA \n")
    search_df = search_value([search_word], df)

    if search_word == "":
        print("koniec")
        break
    else:

        df1 = pd.DataFrame(search_df, columns=['RMA', 'Nazwa urządzenia',
                           'Nr seryjny przyjęty', 'Nr seryjny wydany', 'UWAGI'])
        df1 = df1.replace('nan', '')
        writer.save()
        df2 = pd.read_excel(path_excel_save, sheet_name='arkusz_test')
        df3 = df2.append(df1, ignore_index=True)
        df2 = df2.fillna('')
        df3.to_excel(writer, sheet_name='arkusz_test', index=False)

        df3 = df3.fillna('')
        print(df3)

        writer.save()

        continue

for i in range(df3.shape[0]):
    doc.tables[0].add_row()
    for j in range(df3.shape[-1]):
        table2 = doc.tables[0]
        table2.cell(i+1, j+1).text = str(df3.values[i, j])
        table2.cell(i+1, 0).text = str(i+1)


doc.save("rozliczenie " + cr2 + " " + nazwa+".docx")
del excel_book[df_name]
excel_book.save(path_excel_save)
excel_book.close()




def my_click():
    label=Label(main_window, text="wynik to " + c)
    label.pack()

my_label = Label(main_window, text=ct)
my_label2 = Label(main_window, text=cr2)
my_button = Button(main_window, text="click me", command=my_click)
e1= Entry(main_window, text="a")

e1.pack()

my_label.pack()
my_button.pack()
main_window.mainloop()